#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 기본 스타일 설정 ──────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 페이지 여백 설정
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 헬퍼 함수들 ──────────────────────────────────────────

def set_font(run, size=10, bold=False, color=None, italic=False):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 상단 간격
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(37, 99, 235))  # 파란색
    # 하단 구분선 효과를 위해 border 추가
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '2563EB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(30, 64, 175))
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(55, 65, 81))
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, size=10)
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.8)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, size=10, bold=True)
        r2 = p.add_run(text)
        set_font(r2, size=10)
    else:
        run = p.add_run(text)
        set_font(run, size=10)
    return p

def add_numbered(doc, text, num, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    if bold_prefix:
        r1 = p.add_run(f"{num}. {bold_prefix}")
        set_font(r1, size=10, bold=True)
        r2 = p.add_run(text)
        set_font(r2, size=10)
    else:
        run = p.add_run(f"{num}. {text}")
        set_font(run, size=10)
    return p

def add_callout(doc, text, bg_type='info'):
    # 콜아웃 박스 (테이블 1행 1열로 구현)
    colors = {
        'info':    (219, 234, 254),  # 연파랑
        'warning': (254, 243, 199),  # 연노랑
        'tip':     (220, 252, 231),  # 연초록
        'danger':  (254, 226, 226),  # 연빨강
    }
    text_colors = {
        'info':    (30, 64, 175),
        'warning': (146, 64, 14),
        'tip':     (20, 83, 45),
        'danger':  (153, 27, 27),
    }
    icons = {'info': 'ℹ️  ', 'warning': '⚠️  ', 'tip': '💡  ', 'danger': '🚨  '}

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    cell.width = Inches(6)

    # 배경색
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(*colors[bg_type])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

    # 테두리 제거
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)

    # 셀 패딩
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'left', 'bottom', 'right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '120')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

    p = cell.paragraphs[0]
    run = p.add_run(icons[bg_type] + text)
    set_font(run, size=9.5, color=text_colors[bg_type])
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 헤더 행
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=9.5, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 헤더 배경색 (진파랑)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1D4ED8')
        tcPr.append(shd)

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F8FAFC' if ri % 2 == 0 else 'EFF6FF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            if isinstance(cell_text, tuple):
                # (텍스트, bold) 형태
                run = p.add_run(cell_text[0])
                set_font(run, size=9.5, bold=cell_text[1])
            else:
                run = p.add_run(str(cell_text))
                set_font(run, size=9.5)
            # 행 배경색
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg)
            tcPr.append(shd)

    # 열 너비
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_flow_box(doc, steps):
    """화살표 흐름도를 텍스트로 표현"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    for i, step in enumerate(steps):
        run = p.add_run(step)
        set_font(run, size=9.5, bold=(i % 2 == 0))
        if i < len(steps) - 1:
            arrow = p.add_run('  →  ')
            set_font(arrow, size=9.5, color=(148, 163, 184))
    return p

def add_page_break(doc):
    doc.add_page_break()

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════
#  표지
# ═══════════════════════════════════════════════════════════

# 상단 여백
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Referio')
set_font(run, size=36, bold=True, color=(37, 99, 235))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('운영 매뉴얼')
set_font(run, size=28, bold=True, color=(30, 41, 59))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Operator Guide for Beginners')
set_font(run, size=13, italic=True, color=(100, 116, 139))

for _ in range(3):
    doc.add_paragraph()

# 구분선
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─' * 40)
set_font(run, color=(203, 213, 225))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('대상: 신규 오퍼레이터 (인턴/운영 담당자)')
set_font(run, size=10, color=(71, 85, 105))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('작성일: 2026년 3월 23일  |  버전: 1.0')
set_font(run, size=10, color=(71, 85, 105))

add_page_break(doc)

# ═══════════════════════════════════════════════════════════
#  목차
# ═══════════════════════════════════════════════════════════

add_heading1(doc, '목차')

toc_items = [
    ('1장', 'Referio가 뭔가요? — 서비스 개요', '3'),
    ('2장', '나는 여기서 뭐하는 사람인가? — 오퍼레이터 역할', '4'),
    ('3장', '등장인물 소개 — 광고주, 파트너, 운영자', '5'),
    ('4장', '어필리에이트가 뭔가요? — 개념 이해', '6'),
    ('5장', '전체 흐름 한눈에 보기', '7'),
    ('6장', '광고주 포털 운영 가이드', '8'),
    ('7장', '파트너 포털 운영 가이드', '12'),
    ('8장', '관리자 패널 운영 가이드', '14'),
    ('9장', '이슈 상황별 대응 매뉴얼', '15'),
    ('10장', '용어 사전', '19'),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{num}  ')
    set_font(r1, size=10, bold=True, color=(37, 99, 235))
    r2 = p.add_run(title)
    set_font(r2, size=10)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════
#  1장: 서비스 개요
# ═══════════════════════════════════════════════════════════

add_heading1(doc, '1장  Referio가 뭔가요?')

add_heading2(doc, '한 줄 정의')
add_body(doc, 'Referio는 "소개비 자동 정산 플랫폼"입니다.')
add_body(doc, '기업이 소개자(파트너)에게 고객을 데려와 달라고 부탁하고, 실제로 계약이 되면 소개비를 자동으로 지급하는 시스템입니다.')

add_heading2(doc, '쉬운 비유')
add_body(doc, '일상에서 이런 경험 있으셨나요?')
add_bullet(doc, '"○○ 인테리어 업체 좋더라, 나한테서 소개받았다고 하면 10만원 드려요"')
add_bullet(doc, '"이 앱 친구에게 추천하면 양쪽 다 5,000원 적립"')
add_bullet(doc, '"아는 사람 통해서 보험 들면 사은품 드립니다"')
add_body(doc, 'Referio는 이 "소개비" 시스템을 B2B 기업용으로 디지털화한 플랫폼입니다.')
add_body(doc, '소개 링크 발급 → 고객 추적 → 계약 확인 → 수수료 정산까지 모든 과정을 자동화합니다.')

add_heading2(doc, '왜 필요한가요?')
add_body(doc, '기존 방식의 문제점:')
add_bullet(doc, '누가 소개한 고객인지 일일이 엑셀로 추적')
add_bullet(doc, '소개비 계산을 수동으로 하다 보니 실수 발생')
add_bullet(doc, '파트너(소개자)가 "내가 소개한 거 맞냐"고 다툼 발생')
add_bullet(doc, '파트너가 본인 실적을 확인할 방법이 없어 동기 저하')
add_body(doc, 'Referio의 해결:')
add_bullet(doc, '고유 추천 링크로 누가 소개했는지 100% 자동 추적')
add_bullet(doc, '계약 완료 시 정산 자동 생성 (실수 없음)')
add_bullet(doc, '파트너가 실시간으로 본인 실적 확인 가능')
add_bullet(doc, '광고주는 영업사원 없이 파트너 네트워크로 영업 채널 확대')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════
#  2장: 오퍼레이터 역할
# ═══════════════════════════════════════════════════════════

add_heading1(doc, '2장  나는 여기서 뭐하는 사람인가?')

add_heading2(doc, '오퍼레이터(운영자)의 정의')
add_body(doc, '오퍼레이터는 Referio 플랫폼 전체를 관찰하고, 문제가 생기면 해결하고, 서비스가 잘 굴러가도록 관리하는 사람입니다.')
add_callout(doc, '쉽게 말하면: 놀이공원 안전요원과 같습니다. 놀이기구(서비스)가 잘 작동하는지 보고, 문제가 생기면 빠르게 조치합니다. 직접 놀이기구를 타는 사람(광고주/파트너)이 아니라, 전체를 감시하고 도와주는 사람입니다.', 'tip')

add_heading2(doc, '오퍼레이터가 하는 3가지 핵심 역할')

add_heading3(doc, '① 관찰하기 (Observe)')
add_body(doc, '플랫폼에서 어떤 일이 일어나고 있는지 매일 확인합니다.')
add_bullet(doc, '오늘 신규 리드가 몇 건 들어왔는가?')
add_bullet(doc, '정산 대기 중인 건이 오래 쌓이고 있지는 않은가?')
add_bullet(doc, '파트너 신청이 장시간 승인 대기 중은 아닌가?')
add_bullet(doc, '광고주나 파트너에게서 문의가 들어오지 않았는가?')

add_heading3(doc, '② 이슈 해결하기 (Troubleshoot)')
add_body(doc, '문제가 생겼을 때 원인을 파악하고 해결합니다.')
add_bullet(doc, '"추천 링크가 작동 안 해요" → 원인 찾아 해결')
add_bullet(doc, '"정산이 안 됐어요" → 데이터 확인 후 처리')
add_bullet(doc, '"비밀번호를 잊었어요" → 재설정 도와주기')
add_bullet(doc, '"파트너 승인이 안 되네요" → 광고주에게 확인 요청')

add_heading3(doc, '③ 잘 운영되게 하기 (Maintain)')
add_body(doc, '서비스가 항상 원활하게 돌아가도록 환경을 유지합니다.')
add_bullet(doc, '광고주가 캠페인 설정을 올바르게 했는지 확인')
add_bullet(doc, '신규 광고주 온보딩 지원 (처음 가입한 광고주 안내)')
add_bullet(doc, 'Airtable 연동 등 기술 설정 도와주기')
add_bullet(doc, '이상한 데이터나 오류가 있으면 즉시 보고')

add_heading2(doc, '오퍼레이터가 하지 않아도 되는 것')
add_bullet(doc, '코드를 직접 수정하거나 개발하는 일 (개발팀 몫)')
add_bullet(doc, '광고주 영업 관련 결정 (광고주 본인 몫)')
add_bullet(doc, '파트너와의 수수료 협상 (광고주 몫)')
add_callout(doc, '요약: 오퍼레이터는 "플랫폼이 잘 작동하는지 감시하고, 막히면 뚫어주는 사람"입니다.', 'info')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════
#  3장: 등장인물
# ═══════════════════════════════════════════════════════════

add_heading1(doc, '3장  등장인물 소개')

add_body(doc, 'Referio에는 세 종류의 사용자가 있습니다. 각자의 역할을 명확히 이해해야 이슈 발생 시 누구에게 연락해야 하는지 알 수 있습니다.')

add_heading2(doc, '① 광고주 (Advertiser)')
add_table(doc,
    ['항목', '내용'],
    [
        ['누구?', 'B2B 기업 (소프트웨어 회사, IT 솔루션 업체 등)'],
        ['목표', '파트너 채널을 통해 새로운 고객(리드)을 유치하고 싶은 기업'],
        ['예시', '인사관리 SaaS 회사, ERP 소프트웨어 기업, 마케팅 툴 회사'],
        ['하는 일', '파트너 승인 → 리드 영업 관리 → 계약 완료 처리 → 수수료 지급'],
        ['로그인 방식', '광고주 ID + 사용자 ID + 비밀번호 (3단계)'],
        ['접속 URL', 'referio.kr/advertiser/login'],
    ],
    col_widths=[3.5, 12]
)

add_heading2(doc, '② 파트너 (Partner)')
add_table(doc,
    ['항목', '내용'],
    [
        ['누구?', '추천 링크를 공유해서 수수료를 받는 개인 또는 팀'],
        ['목표', '자신의 채널(블로그, 유튜브 등)로 기업 고객을 소개하고 수수료 수익'],
        ['예시', '비즈니스 블로거, IT 유튜버, 컨설턴트, 에이전시, 지인 네트워크'],
        ['하는 일', '프로그램 가입 → 추천 링크 공유 → 실적 확인 → 수수료 수령'],
        ['로그인 방식', '이메일 + 비밀번호'],
        ['접속 URL', 'referio.kr/login'],
    ],
    col_widths=[3.5, 12]
)

add_heading2(doc, '③ 운영자 (Admin = 우리)')
add_table(doc,
    ['항목', '내용'],
    [
        ['누구?', 'Referio 운영팀 (오퍼레이터 본인 포함)'],
        ['목표', '플랫폼 전체가 원활하게 작동하도록 관리 및 지원'],
        ['하는 일', '전체 현황 모니터링, 이슈 해결, 신규 광고주 온보딩 지원'],
        ['접속 URL', 'referio.kr/admin (외부 공개 X)'],
    ],
    col_widths=[3.5, 12]
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════
#  4장: 어필리에이트
# ═══════════════════════════════════════════════════════════

add_heading1(doc, '4장  어필리에이트(Affiliate)가 뭔가요?')

add_heading2(doc, '어필리에이트 = 제휴 마케팅')
add_body(doc, '어필리에이트(Affiliate)는 영어로 "제휴하다, 연결하다"는 뜻입니다.')
add_body(doc, '어필리에이트 마케팅은 쉽게 말해 "성과 기반 소개비 마케팅"입니다.')
add_callout(doc, '핵심 원리: 광고비를 선불로 쓰지 않고, 실제 계약이 성사됐을 때만 수수료를 지급합니다. → 광고주 입장에서 리스크가 없음!', 'tip')

add_heading2(doc, '일상 속 어필리에이트 사례')
add_table(doc,
    ['사례', '광고주', '파트너', '수수료 방식'],
    [
        ['쿠팡 파트너스', '쿠팡', '블로거/유튜버', '구매 금액의 3%'],
        ['네이버 블로그 광고', '광고 기업', '블로거', '클릭당 / 구매당'],
        ['보험 설계사 소개', '보험사', '지인 네트워크', '계약당 수수료'],
        ['Referio (우리)', 'B2B 기업', 'IT 인플루언서/컨설턴트', '리드당 + 계약당'],
    ],
    col_widths=[4, 3.5, 4, 4]
)

add_heading2(doc, 'Referio 어필리에이트의 특징')
add_bullet(doc, 'B2B 특화: 일반 쇼핑몰이 아닌 기업 소프트웨어/서비스 전문', bold_prefix='')
add_bullet(doc, '2단계 수수료: 리드(문의)가 들어와도 수수료 + 계약이 성사되면 추가 수수료', bold_prefix='')
add_bullet(doc, '파이프라인 관리: 단순 클릭 추적이 아닌 영업 CRM처럼 상태 관리', bold_prefix='')
add_bullet(doc, '자동 정산: 계약 완료 시 수수료 자동 계산 및 정산 생성', bold_prefix='')

add_heading2(doc, '어필리에이트 vs 일반 광고 비교')
add_table(doc,
    ['구분', '일반 광고 (배너/검색 광고)', 'Referio 어필리에이트'],
    [
        ['비용 발생 시점', '광고 노출/클릭 시 (결과 관계없이)', '실제 계약 성사 시'],
        ['리스크', '광고비 낭비 가능', '낭비 없음 (성과 기반)'],
        ['파트너 동기', '없음', '실적 클수록 수수료 증가'],
        ['추적 방식', '쿠키/픽셀', '고유 추천 링크'],
    ],
    col_widths=[4, 5.5, 6]
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════
#  5장: 전체 흐름
# ═══════════════════════════════════════════════════════════

add_heading1(doc, '5장  전체 흐름 한눈에 보기')

add_heading2(doc, '큰 그림: 누가, 뭘 하는가')

add_table(doc,
    ['단계', '누가', '무엇을', '어디서'],
    [
        ['1', '광고주', '가입하고 캠페인(수수료 정책) 설정', 'referio.kr/advertiser'],
        ['2', '광고주', '파트너 모집 (링크 공유 또는 직접 초대)', 'referio.kr/advertiser/partners'],
        ['3', '파트너', '가입하고 프로그램 신청', 'referio.kr/signup'],
        ['4', '광고주', '파트너 신청 승인 → 추천 링크 자동 발급', 'referio.kr/advertiser/partners'],
        ['5', '파트너', '추천 링크를 블로그/SNS/유튜브에 공유', 'Referio 외부'],
        ['6', '잠재고객', '링크 클릭 → 문의 폼 작성', 'referio.kr/inquiry/광고주ID'],
        ['7', 'Referio 시스템', '리드 자동 생성 + 광고주에게 알림', '자동'],
        ['8', '광고주', '리드 영업 관리 (통화, 미팅, 계약 협상)', 'referio.kr/advertiser/referrals'],
        ['9', '광고주', '계약 완료 처리', 'referio.kr/advertiser/referrals'],
        ['10', 'Referio 시스템', '정산 자동 생성', '자동'],
        ['11', '광고주', '파트너에게 실제 송금 후 완료 처리', 'referio.kr/advertiser/settlements'],
        ['12', '파트너', '수수료 수령 확인', 'referio.kr/dashboard/settlements'],
    ],
    col_widths=[1.2, 2.5, 6.3, 5.5]
)

add_heading2(doc, '오퍼레이터는 어디서 무엇을 보는가')
add_body(doc, '오퍼레이터는 위 1~12 단계의 모든 데이터를 관리자 패널에서 볼 수 있습니다.')
add_bullet(doc, '이상한 데이터 → 직접 수정', bold_prefix='관찰: ')
add_bullet(doc, '"링크가 안 된다", "정산이 없다" 등 → 원인 파악 후 해결', bold_prefix='이슈 해결: ')
add_bullet(doc, '오래 방치된 리드, 미승인 파트너 → 광고주에게 알림', bold_prefix='유지보수: ')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════
#  6장: 광고주 포털
# ═══════════════════════════════════════════════════════════

add_heading1(doc, '6장  광고주 포털 운영 가이드')
add_callout(doc, '오퍼레이터가 광고주 계정으로 직접 로그인하는 경우는 드뭅니다. 그러나 광고주가 "어디서 어떻게 해야 하나요?"라고 물을 때 안내해줄 수 있어야 합니다.', 'info')

add_heading2(doc, '6-1. 광고주 메뉴 구성')
add_table(doc,
    ['메뉴 이름', '하는 일', '오퍼레이터가 신경 쓸 때'],
    [
        ['대시보드', '전체 현황 요약 (리드 수, 수익, 파트너 수)', '매일 현황 체크'],
        ['파트너 관리', '파트너 승인/거절, 등급 관리', '신청이 장시간 방치될 때'],
        ['리드 관리', '고객 파이프라인 (영업 상태 관리)', '리드 상태가 안 바뀔 때'],
        ['정산 관리', '수수료 지급 처리', '정산 대기가 쌓일 때'],
        ['캠페인', '수수료 단가·조건 설정', '단가 오류 문의 시'],
        ['프로모션', '파트너 대상 이벤트 공지', '이벤트 관련 문의 시'],
        ['협업', '파트너에게 콘텐츠 제작 의뢰', '협업 상태 이슈 시'],
        ['설정', '비밀번호 변경, Airtable 연동', 'Airtable 연동 문의 시'],
    ],
    col_widths=[3.5, 6, 6]
)

add_heading2(doc, '6-2. 파트너 관리 — 승인 처리')
add_body(doc, '파트너가 프로그램에 신청하면 광고주가 승인해야 합니다. 승인 전까지 파트너는 추천 링크를 받을 수 없습니다.')

add_heading3(doc, '광고주가 신청을 방치할 때')
add_body(doc, '상황: 파트너가 "신청했는데 링크를 못 받았어요"라고 문의할 때')
add_numbered(doc, '관리자 패널에서 해당 파트너의 신청 날짜 확인', 1)
add_numbered(doc, '48시간 이상 대기 중이면 광고주에게 알림 (이메일/슬랙)', 2)
add_numbered(doc, '광고주가 응답 없으면 에스컬레이션 (책임자 보고)', 3)

add_heading3(doc, '승인 시 자동으로 일어나는 일')
add_bullet(doc, 'partner_programs 테이블에 status = approved 업데이트')
add_bullet(doc, '파트너 전용 추천 코드(referral_code) 생성')
add_bullet(doc, '파트너 대시보드에 추천 링크 표시됨')

add_heading2(doc, '6-3. 리드 관리 — 영업 상태 이해')
add_body(doc, '리드(Lead) = 파트너 링크를 통해 들어온 잠재 고객 문의 1건')
add_body(doc, '각 리드는 아래 상태 중 하나를 가집니다:')

add_table(doc,
    ['상태', '의미', '다음 단계'],
    [
        ['pending', '방금 들어온 신규 문의', '광고주가 연락을 시도해야 함'],
        ['call_1', '1차 연락 완료', '추가 미팅/데모 진행'],
        ['call_2', '2차 연락 완료', '계약 조건 협상'],
        ['call_3', '3차 연락 완료', '계약서 작성'],
        ['completed', '계약 완료!', '정산 자동 생성됨'],
        ['invalid', '무효 처리', '연락 안 됨, 조건 불일치 등'],
        ['duplicate', '중복 리드', '이미 알고 있던 고객'],
    ],
    col_widths=[2.5, 5.5, 7.5]
)

add_callout(doc, '중요: 리드가 "completed" + "is_valid = true" 상태가 되면 정산이 자동 생성됩니다. 이 두 조건을 모두 충족해야 합니다.', 'warning')

add_heading2(doc, '6-4. 정산 관리 — 수수료 지급 흐름')

add_heading3(doc, '정산 자동 생성 조건')
add_table(doc,
    ['정산 유형', '생성 조건', '금액 기준'],
    [
        ['유효 리드 수수료', 'is_valid = true로 설정 시', '캠페인의 유효 리드 단가'],
        ['계약 수수료', 'contract_status = completed 설정 시', '캠페인의 계약 단가 또는 %'],
    ],
    col_widths=[4.5, 5.5, 5.5]
)

add_heading3(doc, '정산 처리 순서')
add_numbered(doc, '광고주가 파트너에게 실제 계좌이체로 송금', 1)
add_numbered(doc, 'Referio에서 "완료 처리" 버튼 클릭', 2)
add_numbered(doc, '파트너 대시보드에 "완료" 상태로 표시됨', 3)
add_callout(doc, '오퍼레이터 주의: 정산은 Referio가 직접 돈을 보내는 것이 아닙니다. 광고주가 파트너에게 송금하고, Referio에서 "완료" 처리를 누르는 구조입니다.', 'warning')

add_heading2(doc, '6-5. 캠페인 설정 — 수수료 정책')
add_body(doc, '캠페인은 광고주가 파트너에게 지급할 수수료 규칙을 정의합니다.')
add_table(doc,
    ['설정 항목', '설명', '예시'],
    [
        ['유효 리드 단가', '문의 1건당 지급액', '30,000원'],
        ['계약 수수료', '계약 1건당 지급액 또는 %', '100,000원 또는 5%'],
        ['중복 체크 기간', '같은 연락처가 재유입 방지되는 기간', '30일'],
        ['유효 리드 기한', '문의 후 "유효" 판단까지 목표일', '7일'],
        ['계약 기한', '계약 완료까지 목표일', '30일'],
    ],
    col_widths=[4, 6, 5.5]
)

add_heading2(doc, '6-6. Airtable 연동 설정 (4단계)')
add_body(doc, 'Airtable(외부 스프레드시트 도구)에서 리드가 자동으로 Referio로 들어오게 하는 기능입니다.')
add_numbered(doc, 'Airtable API Key 입력 → 연결 확인', 1)
add_numbered(doc, 'Airtable 컬럼 ↔ Referio 필드 연결 (이름, 전화번호 등)', 2)
add_numbered(doc, 'Webhook URL 복사 (Airtable에 붙여넣을 코드)', 3)
add_numbered(doc, 'Airtable에서 자동화(Automation) 설정', 4)
add_callout(doc, 'Airtable 연동이 안 될 때는 9장 이슈 대응 매뉴얼을 참고하세요.', 'info')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════
#  7장: 파트너 포털
# ═══════════════════════════════════════════════════════════

add_heading1(doc, '7장  파트너 포털 운영 가이드')

add_heading2(doc, '7-1. 파트너 가입 프로세스')
add_body(doc, '파트너가 처음 가입할 때의 전체 단계입니다:')
add_table(doc,
    ['단계', '파트너가 할 일', '결과'],
    [
        ['1', 'referio.kr/signup 접속', '가입 페이지 열림'],
        ['2', '이메일 + 비밀번호 입력', '인증 이메일 발송'],
        ['3', '이메일함에서 인증 링크 클릭', '계정 활성화'],
        ['4', '대시보드에서 프로그램 탐색', '광고주 목록 표시'],
        ['5', '원하는 프로그램 신청', '광고주에게 신청 알림'],
        ['6', '광고주 승인 대기', '승인되면 알림'],
        ['7', '대시보드에서 추천 링크 확인', '링크 공유 시작 가능'],
    ],
    col_widths=[1.2, 7, 7]
)

add_heading2(doc, '7-2. 브랜디드 가입 링크')
add_body(doc, '광고주가 자신만의 브랜드로 파트너를 모집하는 전용 가입 페이지를 만들 수 있습니다.')
add_body(doc, '형식:  referio.kr/signup/{광고주ID}')
add_body(doc, '예시:  referio.kr/signup/puzl_co_kr')
add_callout(doc, '브랜디드 가입 링크로 들어온 파트너는 해당 광고주의 프로그램에 자동 연결됩니다. 별도 신청이 필요 없을 수 있습니다.', 'tip')

add_heading2(doc, '7-3. 파트너 대시보드 주요 항목')
add_table(doc,
    ['항목', '설명', '오퍼레이터 체크 포인트'],
    [
        ['추천 URL', '파트너가 공유할 고유 링크', '링크 형식이 올바른지 확인'],
        ['실적 통계', '클릭 수, 리드 수, 계약 수, 총 수익', '수치가 이상하면 DB 확인'],
        ['티어 배지', '현재 파트너 등급 (Authorized/Silver/Gold/Platinum)', '등급 오류 문의 시 DB 확인'],
        ['마일스톤', '달성 목표 및 진행률', '표시 오류 시 보고'],
        ['정산 현황', '수수료 내역 및 지급 상태', '미지급 장기 방치 확인'],
    ],
    col_widths=[3, 5.5, 7]
)

add_heading2(doc, '7-4. 추천 링크 형식')
add_body(doc, '파트너에게 발급되는 추천 링크는 두 가지 형태입니다:')
add_table(doc,
    ['형태', 'URL 형식', '설명'],
    [
        ['직접 문의 폼', 'referio.kr/inquiry/{광고주ID}?ref={코드}', '문의 폼으로 바로 연결'],
        ['보안 리다이렉트', 'referio.kr/security?ref={코드}', '자동으로 문의 폼으로 이동'],
    ],
    col_widths=[3.5, 8, 4]
)
add_callout(doc, '두 링크 모두 동일하게 작동합니다. 파트너 대시보드에서 복사한 링크를 그대로 사용하면 됩니다.', 'tip')

add_heading2(doc, '7-5. 파트너 등급(티어) 시스템')
add_table(doc,
    ['등급', '의미', '수수료'],
    [
        ['Authorized', '기본 등급 (승인된 파트너)', '기본 단가'],
        ['Silver', '실적 쌓인 파트너', '기본 단가 + α'],
        ['Gold', '우수 파트너', '높은 단가'],
        ['Platinum', '최우수 파트너', '최고 단가'],
    ],
    col_widths=[3.5, 7, 5]
)
add_callout(doc, '현재 티어는 광고주가 수동으로 변경합니다. 자동 승급 기능은 아직 개발 중입니다.', 'info')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════
#  8장: 관리자 패널
# ═══════════════════════════════════════════════════════════

add_heading1(doc, '8장  관리자 패널 운영 가이드')
add_callout(doc, '관리자 패널은 Referio 운영팀만 접근할 수 있습니다. URL: referio.kr/admin', 'warning')

add_heading2(doc, '8-1. 메뉴 구성')
add_table(doc,
    ['메뉴', '볼 수 있는 것', '언제 사용'],
    [
        ['대시보드', '전체 플랫폼 현황', '매일 아침 현황 체크'],
        ['리드 관리', '모든 광고주의 전체 리드 데이터', '특정 리드 이슈 조사'],
        ['정산 관리', '모든 정산 내역', '정산 미지급 이슈 조사'],
        ['파트너 관리', '모든 파트너 계정', '파트너 계정 이슈'],
        ['광고주 관리', '모든 광고주 계정', '광고주 계정 이슈'],
        ['캠페인 관리', '모든 캠페인 설정', '수수료 설정 오류 확인'],
    ],
    col_widths=[3.5, 6, 6]
)

add_heading2(doc, '8-2. 일일 체크리스트')
add_body(doc, '매일 아침 관리자 패널에서 확인해야 할 항목:')
add_bullet(doc, '신규 리드가 정상적으로 들어오고 있는가?')
add_bullet(doc, '48시간 이상 대기 중인 파트너 신청이 있는가?')
add_bullet(doc, '1주일 이상 처리되지 않은 정산 대기 건이 있는가?')
add_bullet(doc, '오류로 인해 status가 NULL이거나 비정상인 데이터가 있는가?')
add_bullet(doc, '광고주/파트너로부터 들어온 문의가 있는가?')

add_heading2(doc, '8-3. Supabase 직접 접근 (고급)')
add_body(doc, '일부 이슈는 관리자 패널에서 해결이 안 되어 데이터베이스에 직접 접근해야 합니다.')
add_callout(doc, '주의: 데이터베이스를 직접 수정하는 것은 위험합니다. 반드시 책임자 확인 후 진행하세요. 삭제 작업은 복구가 어렵습니다!', 'danger')

add_body(doc, 'Supabase 접근 방법:')
add_numbered(doc, 'supabase.com에 접속하여 로그인', 1)
add_numbered(doc, '프로젝트 "eqdnirtgmevhobmycxzn" 선택', 2)
add_numbered(doc, 'Table Editor에서 원하는 테이블 검색', 3)
add_numbered(doc, 'SQL Editor에서 쿼리로 데이터 조회/수정', 4)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════
#  9장: 이슈 대응
# ═══════════════════════════════════════════════════════════

add_heading1(doc, '9장  이슈 상황별 대응 매뉴얼')
add_body(doc, '실제로 자주 발생하는 이슈들과 해결 방법입니다. 상황을 먼저 파악하고, 순서대로 체크해보세요.')

add_heading2(doc, '이슈 #1  "파트너 링크를 눌렀는데 오류가 나요"')
add_heading3(doc, '원인 파악 체크리스트')
add_numbered(doc, '링크 URL 형식 확인: ?ref=코드 파라미터가 있는가?', 1)
add_numbered(doc, '관리자 패널 → 파트너 관리에서 해당 파트너 검색', 2)
add_numbered(doc, 'partner_programs.status = approved 인지 확인', 3)
add_numbered(doc, '해당 referral_code가 존재하는지 확인', 4)
add_heading3(doc, '해결 방법')
add_bullet(doc, '미승인 상태: 광고주에게 승인 요청')
add_bullet(doc, '코드 없음: Supabase에서 referral_code 직접 확인 후 수정')
add_bullet(doc, '링크 형식 오류: 파트너에게 올바른 링크 재전달')

add_divider(doc)
add_heading2(doc, '이슈 #2  "리드를 넣었는데 광고주 대시보드에 안 보여요"')
add_heading3(doc, '원인 파악 체크리스트')
add_numbered(doc, '파트너가 올바른 광고주 URL로 문의했는지 확인', 1)
add_body(doc, '문의 폼 URL: referio.kr/inquiry/{광고주ID}?ref={추천코드}', indent=True)
add_numbered(doc, 'Supabase → referrals 테이블에서 데이터 존재 여부 확인', 2)
add_numbered(doc, '광고주가 올바른 계정으로 로그인했는지 확인', 3)
add_heading3(doc, '해결 방법')
add_bullet(doc, 'referrals 테이블에 데이터 없음: 링크/코드 오류. 파트너에게 올바른 링크 재제공')
add_bullet(doc, 'advertiser_id 불일치: Supabase에서 데이터 직접 수정')
add_bullet(doc, '데이터는 있는데 안 보임: 광고주 로그인 계정 확인')

add_divider(doc)
add_heading2(doc, '이슈 #3  "계약 완료로 바꿨는데 정산이 자동 생성이 안 돼요"')
add_heading3(doc, '원인 파악 체크리스트')
add_numbered(doc, 'Supabase → referrals 테이블에서 해당 리드 확인', 1)
add_numbered(doc, 'is_valid 컬럼이 true인지 확인 (false나 null이면 안 생김)', 2)
add_numbered(doc, 'partner_id 컬럼이 NULL이 아닌지 확인', 3)
add_numbered(doc, 'Supabase → settlements 테이블에서 이미 생성됐는지 확인', 4)
add_heading3(doc, '해결 방법')
add_bullet(doc, 'is_valid가 false/null: 광고주에게 유효 처리 먼저 하도록 안내')
add_bullet(doc, 'partner_id가 NULL: Supabase에서 직접 partner_id 입력 후 책임자 보고')
add_bullet(doc, '정산이 이미 있음: 중복 생성 방지 로직 정상 작동. 기존 정산 확인')
add_callout(doc, '긴급 처리가 필요할 경우: Supabase SQL Editor에서 settlements 테이블에 직접 INSERT 가능. 반드시 책임자 승인 후 진행.', 'danger')

add_divider(doc)
add_heading2(doc, '이슈 #4  "광고주가 로그인이 안 된다고 해요"')
add_heading3(doc, '원인 파악 체크리스트')
add_numbered(doc, '입력 방식 확인: 광고주 ID → 사용자 ID → 비밀번호 (3단계)', 1)
add_body(doc, '광고주 ID 예시: puzl_co_kr   /   사용자 ID 예시: admin', indent=True)
add_numbered(doc, 'Supabase → advertisers 또는 advertiser_users 테이블에서 계정 존재 여부 확인', 2)
add_numbered(doc, 'advertiser_sessions 테이블에서 만료 세션 확인 (세션 유효기간 7일)', 3)
add_heading3(doc, '해결 방법')
add_bullet(doc, '입력 오류: 올바른 로그인 형식 안내')
add_bullet(doc, '계정 없음: 가입 여부 확인 후 재가입 안내')
add_bullet(doc, '비밀번호 분실: advertiser/settings에서 변경 가능. 본인 접속이 안 되면 Supabase에서 bcrypt 해시로 재설정')
add_callout(doc, '광고주 비밀번호를 Supabase에서 초기화할 때는 반드시 책임자 확인 후 진행하세요.', 'danger')

add_divider(doc)
add_heading2(doc, '이슈 #5  "파트너가 비밀번호를 잊어버렸어요"')
add_heading3(doc, '해결 방법 (간단)')
add_numbered(doc, 'referio.kr/reset-password 링크를 파트너에게 전달', 1)
add_numbered(doc, '파트너가 이메일 입력 → 재설정 링크 수신', 2)
add_numbered(doc, '새 비밀번호로 로그인', 3)
add_body(doc, '위 방법이 안 될 경우:')
add_bullet(doc, 'Supabase → Authentication → Users에서 해당 이메일 계정 찾기')
add_bullet(doc, '"Send password recovery" 버튼 클릭')

add_divider(doc)
add_heading2(doc, '이슈 #6  "Airtable에서 리드가 자동으로 안 들어와요"')
add_heading3(doc, '체크리스트')
add_numbered(doc, 'Settings → Airtable 탭에서 API Key 연결 상태 확인', 1)
add_numbered(doc, 'Airtable Automation이 "켜짐" 상태인지 확인', 2)
add_numbered(doc, 'Webhook URL 확인: referio.kr/api/webhook/airtable', 3)
add_numbered(doc, 'X-API-Key 헤더가 포함됐는지 확인', 4)
add_numbered(doc, 'Vercel 로그에서 에러 메시지 확인', 5)
add_callout(doc, 'Airtable 연동 이슈는 대부분 API Key 오류 또는 Automation 설정 오류입니다. 먼저 광고주에게 재연결을 시도하도록 안내하세요.', 'info')

add_divider(doc)
add_heading2(doc, '이슈 #7  "파트너 신청을 했는데 승인이 안 돼요"')
add_heading3(doc, '원인 및 해결')
add_numbered(doc, '관리자 패널에서 해당 파트너의 partner_programs.status 확인', 1)
add_numbered(doc, 'status = pending이면 → 광고주에게 직접 연락하여 승인 요청', 2)
add_numbered(doc, '48시간 이상 방치 → 광고주에게 알림 발송', 3)
add_numbered(doc, '광고주가 지속 미응답 → 책임자 보고', 4)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════
#  10장: 용어 사전
# ═══════════════════════════════════════════════════════════

add_heading1(doc, '10장  용어 사전')
add_body(doc, '처음 보면 생소한 단어들을 쉽게 설명합니다.')

add_table(doc,
    ['용어', '쉬운 설명'],
    [
        ['광고주 (Advertiser)', '파트너에게 "고객 데려오면 수수료 줄게요"라고 하는 기업'],
        ['파트너 (Partner)', '추천 링크를 공유해서 수수료를 받는 개인 또는 팀'],
        ['리드 (Lead/Referral)', '파트너 링크를 통해 들어온 잠재 고객 문의 1건'],
        ['추천 코드 (Referral Code)', '파트너별 고유 식별 코드. 링크에 ?ref=코드 형태로 붙음'],
        ['추천 링크 (Referral URL)', '파트너가 SNS/블로그에 공유하는 고유 링크'],
        ['정산 (Settlement)', '파트너에게 지급하는 수수료 내역'],
        ['캠페인', '광고주가 설정하는 수수료 단가·조건 규칙'],
        ['티어 (Tier)', '파트너 등급. Authorized < Silver < Gold < Platinum'],
        ['프로그램 (Program)', '특정 광고주의 파트너 모집 프로그램. 파트너가 가입 신청'],
        ['문의 폼 (Inquiry Form)', '잠재 고객이 이름/연락처/문의를 입력하는 페이지'],
        ['어필리에이트 (Affiliate)', '소개비 기반 마케팅. 성과가 생길 때만 수수료 지급'],
        ['is_valid', '광고주가 리드를 "유효한 잠재 고객"으로 확인했다는 표시 (true/false)'],
        ['contract_status', '리드의 영업 진행 단계 (pending → call_1 → completed 등)'],
        ['광고주 ID', '광고주 계정을 식별하는 텍스트 (예: puzl_co_kr). 로그인 1단계에 사용'],
        ['Supabase', 'Referio가 사용하는 데이터베이스 플랫폼 (PostgreSQL 기반)'],
        ['Vercel', 'Referio 서버가 배포되어 있는 클라우드 플랫폼'],
        ['Airtable', '외부 스프레드시트 도구. Referio와 연동해 리드 자동 유입 가능'],
        ['멀티테넌트', '여러 광고주가 같은 서비스를 사용하지만 데이터는 완전히 분리되는 구조'],
        ['크론 (Cron)', '정해진 시간마다 자동으로 실행되는 작업. Airtable 동기화에 사용'],
        ['웹훅 (Webhook)', '외부 시스템이 Referio에 데이터를 자동 전송하는 방식'],
    ],
    col_widths=[5, 10.5]
)

# ═══════════════════════════════════════════════════════════
#  부록
# ═══════════════════════════════════════════════════════════

add_page_break(doc)
add_heading1(doc, '부록  빠른 참조')

add_heading2(doc, 'A. 주요 URL 모음')
add_table(doc,
    ['용도', 'URL'],
    [
        ['서비스 메인', 'referio.kr'],
        ['파트너 로그인', 'referio.kr/login'],
        ['파트너 가입', 'referio.kr/signup'],
        ['광고주 로그인', 'referio.kr/advertiser/login'],
        ['광고주 가입', 'referio.kr/advertiser/signup'],
        ['관리자 패널', 'referio.kr/admin'],
        ['파트너 비밀번호 재설정', 'referio.kr/reset-password'],
        ['Airtable Webhook 수신', 'referio.kr/api/webhook/airtable'],
    ],
    col_widths=[5, 10.5]
)

add_heading2(doc, 'B. 시스템 접근 정보')
add_table(doc,
    ['시스템', '접근 방법', '용도'],
    [
        ['Supabase', 'supabase.com → 프로젝트 eqdnirtgmevhobmycxzn', '데이터베이스 직접 조회/수정'],
        ['Vercel', 'vercel.com → referio-platform', '서버 로그, 배포 현황'],
        ['GitHub', 'github.com/cse7600/referio-platform', '소스코드 (개발팀 사용)'],
    ],
    col_widths=[3, 6.5, 6]
)
add_callout(doc, '위 시스템에 대한 계정 정보는 별도로 인수인계 받으세요. 이 문서에는 기재하지 않습니다.', 'warning')

add_heading2(doc, 'C. 긴급 상황 에스컬레이션 기준')
add_table(doc,
    ['상황', '대응'],
    [
        ['리드가 3시간 이상 전혀 들어오지 않음', '즉시 책임자 보고'],
        ['정산 데이터 이상 (금액 오류, 중복 생성 등)', '즉시 책임자 보고 후 Supabase 확인'],
        ['광고주가 "데이터가 사라졌다"고 함', '즉시 책임자 보고 (절대 혼자 처리하지 말 것)'],
        ['48시간 이상 미승인 파트너 신청 누적', '광고주에게 알림 발송'],
        ['Airtable 연동 중단 (24시간 이상)', '광고주에게 안내 후 재연동 지원'],
    ],
    col_widths=[8, 7.5]
)

# 최종 메모
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('이 문서에 오류가 있거나 내용 추가가 필요하면 운영팀에 알려주세요.')
set_font(run, size=9, italic=True, color=(148, 163, 184))

# ── 저장 ──────────────────────────────────────────────────
output_path = '/Users/hokang2father/projects/referio-platform/docs/Referio_운영매뉴얼_v1.0.docx'
doc.save(output_path)
print(f"✅ 저장 완료: {output_path}")
